"""
Text loading module for extracting text from various file formats.
Supports: TXT, PDF, DOCX, RTF, HTML, and raw text strings.

Also provides text chunking utilities for handling large documents
that may not fit in an LLM's context window.
"""

import os
from typing import Union, Optional
from abc import ABC, abstractmethod
from pathlib import Path


class TextLoader(ABC):
    """Base class for text loaders."""
    
    @abstractmethod
    def load(self, source: Union[str, Path]) -> str:
        """
        Load text from the given source.
        
        Args:
            source: File path or raw text string.
            
        Returns:
            Extracted text content.
            
        Raises:
            ValueError: If the source is invalid or cannot be processed.
        """
        pass


class TXTLoader(TextLoader):
    """Loader for plain text files."""
    
    def load(self, source: Union[str, Path]) -> str:
        """Load text from a .txt file."""
        source = Path(source)
        if not source.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        try:
            with open(source, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'iso-8859-1', 'cp1252']:
                try:
                    with open(source, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file with standard encodings: {source}")


class PDFLoader(TextLoader):
    """Loader for PDF files."""
    
    def load(self, source: Union[str, Path]) -> str:
        """Load text from a PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF support. "
                "Install it with: pip install PyPDF2"
            )
        
        source = Path(source)
        if not source.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        text_content = []
        try:
            with open(source, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
        except Exception as e:
            raise ValueError(f"Error reading PDF file {source}: {str(e)}")
        
        return '\n'.join(text_content)


class DOCXLoader(TextLoader):
    """Loader for DOCX files."""
    
    def load(self, source: Union[str, Path]) -> str:
        """Load text from a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. "
                "Install it with: pip install python-docx"
            )
        
        source = Path(source)
        if not source.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        try:
            doc = Document(source)
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            return '\n'.join(text_content)
        except Exception as e:
            raise ValueError(f"Error reading DOCX file {source}: {str(e)}")


class RTFLoader(TextLoader):
    """Loader for RTF files."""
    
    def load(self, source: Union[str, Path]) -> str:
        """Load text from an RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError(
                "striprtf is required for RTF support. "
                "Install it with: pip install striprtf"
            )
        
        source = Path(source)
        if not source.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        try:
            with open(source, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            return rtf_to_text(rtf_content)
        except Exception as e:
            raise ValueError(f"Error reading RTF file {source}: {str(e)}")


class HTMLLoader(TextLoader):
    """Loader for HTML files."""
    
    def load(self, source: Union[str, Path]) -> str:
        """Load text from an HTML file."""
        try:
            from html.parser import HTMLParser
            from io import StringIO
        except ImportError:
            raise ImportError("html parser is required for HTML support (built-in module).")
        
        source = Path(source)
        if not source.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        try:
            with open(source, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Simple HTML parser to extract text
            class MLStripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.reset()
                    self.strict = False
                    self.convert_charrefs = True
                    self.text = StringIO()
                
                def handle_data(self, data):
                    self.text.write(data)
                
                def get_data(self):
                    return self.text.getvalue()
            
            stripper = MLStripper()
            stripper.feed(html_content)
            return stripper.get_data()
        except Exception as e:
            raise ValueError(f"Error reading HTML file {source}: {str(e)}")


class RawTextLoader(TextLoader):
    """Loader for raw text strings."""
    
    def load(self, source: Union[str, Path]) -> str:
        """Return the source as-is if it's not a file."""
        if isinstance(source, str):
            return source
        return str(source)


class UniversalTextLoader:
    """
    Universal text loader that automatically detects file type and loads content.
    Handles multiple formats: TXT, PDF, DOCX, RTF, HTML, and raw text strings.
    """
    
    def __init__(self, enable_logging: bool = False):
        """
        Initialize the universal text loader.
        
        Args:
            enable_logging: Whether to enable logging of loaded content info.
        """
        self.logging = enable_logging
        self.loaders = {
            '.txt': TXTLoader(),
            '.pdf': PDFLoader(),
            '.docx': DOCXLoader(),
            '.doc': DOCXLoader(),  # Try DOCX loader for .doc files
            '.rtf': RTFLoader(),
            '.html': HTMLLoader(),
            '.htm': HTMLLoader(),
        }
    
    def load(self, source: Union[str, Path], file_type: Optional[str] = None) -> str:
        """
        Load text from various sources.
        
        Args:
            source: File path (str or Path), or raw text string.
            file_type: Optional file extension (e.g., '.pdf', '.txt'). 
                      If not provided, will be auto-detected from the source.
        
        Returns:
            Extracted text content.
            
        Raises:
            ValueError: If the file type is not supported or content cannot be extracted.
            FileNotFoundError: If the specified file does not exist.
        """
        # Try to detect if source is a file path or raw text
        source_path = Path(source) if not isinstance(source, Path) else source
        
        is_file = source_path.is_file() if isinstance(source_path, Path) else isinstance(source, str) and os.path.isfile(str(source))
        
        if not is_file:
            # Source is raw text, not a file
            if self.logging:
                print("UniversalTextLoader: INFO :: Treating input as raw text string")
            return RawTextLoader().load(source)
        
        # Source is a file path
        source_path = Path(source)
        
        # Determine file type
        if file_type is None:
            file_type = source_path.suffix.lower()
        elif not file_type.startswith('.'):
            file_type = '.' + file_type.lower()
        
        # Get appropriate loader
        if file_type not in self.loaders:
            supported_types = ', '.join(self.loaders.keys())
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {supported_types}"
            )
        
        loader = self.loaders[file_type]
        
        if self.logging:
            print(f"UniversalTextLoader: INFO :: Loading text from {file_type} file: {source_path.name}")
        
        try:
            text = loader.load(source_path)
            if self.logging:
                word_count = len(text.split())
                char_count = len(text)
                print(f"UniversalTextLoader: INFO :: Successfully loaded {word_count} words ({char_count} characters)")
            return text
        except Exception as e:
            raise ValueError(f"Error loading text from {source_path}: {str(e)}")
    
    def supports_file_type(self, file_type: str) -> bool:
        """
        Check if a file type is supported.
        
        Args:
            file_type: File extension (e.g., '.pdf', 'pdf').
            
        Returns:
            True if the file type is supported, False otherwise.
        """
        if not file_type.startswith('.'):
            file_type = '.' + file_type.lower()
        return file_type in self.loaders
    
    @property
    def supported_formats(self) -> list:
        """Get list of supported file formats."""
        return list(self.loaders.keys())


